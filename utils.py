"""
KoKoFish — Utility functions.

Handles ffmpeg detection, file reading (txt/pdf/docx), audio export,
document export, and system monitoring helpers.
"""

import os
import re
import subprocess
import sys
import shutil
import logging
import threading

import psutil
import pdfplumber
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from lang import t

logger = logging.getLogger("KoKoFish.utils")

# ---------------------------------------------------------------------------
# FFmpeg setup
# ---------------------------------------------------------------------------

def get_app_dir() -> str:
    """Return the directory where the application lives."""
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _download_ffmpeg(bin_dir: str, on_progress=None) -> bool:
    """
    Download a pre-built FFmpeg Windows binary from gyan.dev and extract
    ffmpeg.exe + ffprobe.exe into bin_dir.

    Returns True on success, False on failure.
    """
    import urllib.request
    import zipfile
    import tempfile

    # Essentials build: ~75 MB, contains ffmpeg.exe and ffprobe.exe
    URL = "https://www.gyan.dev/ffmpeg/builds/ffmpeg-release-essentials.zip"

    os.makedirs(bin_dir, exist_ok=True)

    try:
        if on_progress:
            on_progress(t("UTILS_DOWNLOADING_FFMPEG"), 0.05)

        logger.info("Downloading FFmpeg from %s", URL)

        tmp_zip = os.path.join(tempfile.gettempdir(), "ffmpeg_download.zip")

        def _reporthook(block_num, block_size, total_size):
            if on_progress and total_size > 0:
                frac = min(block_num * block_size / total_size, 1.0)
                on_progress(f"Downloading FFmpeg... {int(frac * 100)}%", frac * 0.8)

        urllib.request.urlretrieve(URL, tmp_zip, reporthook=_reporthook)

        if on_progress:
            on_progress(t("UTILS_EXTRACTING_FFMPEG"), 0.85)

        with zipfile.ZipFile(tmp_zip, "r") as zf:
            for member in zf.namelist():
                filename = os.path.basename(member)
                if filename in ("ffmpeg.exe", "ffprobe.exe"):
                    source = zf.open(member)
                    dest_path = os.path.join(bin_dir, filename)
                    with open(dest_path, "wb") as dest:
                        dest.write(source.read())
                    logger.info("Extracted %s -> %s", filename, dest_path)

        try:
            os.remove(tmp_zip)
        except OSError:
            pass

        if on_progress:
            on_progress(t("UTILS_FFMPEG_READY"), 1.0)

        return os.path.isfile(os.path.join(bin_dir, "ffmpeg.exe"))

    except Exception as exc:
        logger.error("FFmpeg auto-download failed: %s", exc)
        return False


def setup_ffmpeg(on_progress=None) -> bool:
    """
    Locate FFmpeg and ensure it is available on the OS PATH so libraries
    like pydub can find it. If FFmpeg is not present, automatically
    downloads the official Windows build from gyan.dev into bin/.

    Accepts an optional on_progress(message, fraction) callback for
    splash screen integration.
    """
    app_dir = get_app_dir()
    bin_dir = os.path.join(app_dir, "bin")
    bundled_ffmpeg = os.path.join(bin_dir, "ffmpeg.exe")

    def _activate(ffmpeg_path: str) -> bool:
        """Put the containing folder on PATH and configure pydub."""
        folder = os.path.dirname(ffmpeg_path)
        current_path = os.environ.get("PATH", "")
        if folder not in current_path:
            os.environ["PATH"] = folder + os.pathsep + current_path
        try:
            from pydub import AudioSegment
            AudioSegment.converter = ffmpeg_path
        except ImportError:
            pass
        return True

    # 1. Bundled bin/ folder
    if os.path.isfile(bundled_ffmpeg):
        logger.info("FFmpeg found (bundled): %s", bundled_ffmpeg)
        return _activate(bundled_ffmpeg)

    # 2. System PATH
    path_ffmpeg = shutil.which("ffmpeg")
    if path_ffmpeg:
        logger.info("FFmpeg found (PATH): %s", path_ffmpeg)
        return _activate(path_ffmpeg)

    # 3. Auto-download
    logger.warning("FFmpeg not found -- attempting auto-download.")
    success = _download_ffmpeg(bin_dir, on_progress=on_progress)
    if success:
        logger.info("FFmpeg auto-download succeeded.")
        return _activate(bundled_ffmpeg)

    logger.error("FFmpeg could not be installed automatically.")
    return False




def is_ffmpeg_available() -> bool:
    """Quick check without mutating state."""
    bundled = os.path.join(get_app_dir(), "bin", "ffmpeg.exe")
    return os.path.isfile(bundled) or shutil.which("ffmpeg") is not None


# ---------------------------------------------------------------------------
# Python dependency auto-install
# ---------------------------------------------------------------------------

# Packages that ship separately from the main requirements but are needed at
# runtime.  Each entry: (import_name, pip_name, version_spec)
_RUNTIME_DEPS = [
    ("noisereduce", "noisereduce", ""),   # voice clone denoising
    ("scipy",       "scipy",       ""),   # high-pass filter for reference audio
]


def setup_python_deps(on_progress=None) -> None:
    """
    Ensure optional runtime packages are installed.
    Runs pip install only when the package cannot be imported.
    Non-fatal — logs warnings on failure.
    """
    import importlib
    import subprocess

    venv_python = sys.executable

    for import_name, pip_name, version_spec in _RUNTIME_DEPS:
        try:
            importlib.import_module(import_name)
            continue  # already installed
        except ImportError:
            pass

        pkg = pip_name + version_spec
        logger.info("Auto-installing missing dependency: %s", pkg)
        if on_progress:
            on_progress(t("UTILS_INSTALLING_PKG", pip_name=pip_name), 0.0)
        try:
            result = subprocess.run(
                [venv_python, "-m", "pip", "install", pkg, "--quiet"],
                capture_output=True, text=True, timeout=180,
            )
            if result.returncode == 0:
                logger.info("Installed %s successfully.", pkg)
            else:
                logger.warning(
                    "Failed to install %s: %s",
                    pkg, result.stderr[-300:] if result.stderr else "unknown error",
                )
        except Exception as exc:
            logger.warning("Auto-install of %s failed (non-fatal): %s", pkg, exc)


# ---------------------------------------------------------------------------
# Fish-Speech auto-setup
# ---------------------------------------------------------------------------

def is_kokoro_ready() -> bool:
    """Return True if both Kokoro ONNX model files are present."""
    app_dir = os.path.dirname(os.path.abspath(__file__))
    model_dir = os.path.join(app_dir, "kokoro_models")
    return (
        os.path.isfile(os.path.join(model_dir, "kokoro-v1.0.int8.onnx")) and
        os.path.isfile(os.path.join(model_dir, "voices-v1.0.bin"))
    )


def setup_kokoro(on_progress=None) -> bool:
    """
    Download Kokoro ONNX model files from HuggingFace if missing.

    Source repo : speaches-ai/Kokoro-82M-v1.0-ONNX-int8
    Files       : model.onnx  (~100 MB) → saved as kokoro-v1.0.int8.onnx
                  voices.bin  (~20 MB)  → saved as voices-v1.0.bin

    on_progress(message, fraction) is called throughout.
    Returns True if models are ready, False on failure.
    """
    app_dir = os.path.dirname(os.path.abspath(__file__))
    model_dir = os.path.join(app_dir, "kokoro_models")
    os.makedirs(model_dir, exist_ok=True)

    # (local_filename, filename_in_repo)
    files = [
        ("kokoro-v1.0.int8.onnx", "model.onnx"),
        ("voices-v1.0.bin",        "voices.bin"),
    ]

    try:
        from huggingface_hub import hf_hub_download
    except ImportError:
        logger.error("huggingface_hub not installed; cannot download Kokoro models.")
        return False

    for i, (local_name, repo_path) in enumerate(files):
        dest = os.path.join(model_dir, local_name)
        if os.path.isfile(dest):
            continue
        frac_base = i / len(files)
        if on_progress:
            on_progress(f"Downloading Kokoro {local_name}…", frac_base + 0.05)
        try:
            logger.info("Downloading Kokoro model: %s from speaches-ai/Kokoro-82M-v1.0-ONNX-int8", repo_path)
            path = hf_hub_download(
                repo_id="speaches-ai/Kokoro-82M-v1.0-ONNX-int8",
                filename=repo_path,
                local_dir=model_dir,
                local_dir_use_symlinks=False,
            )
            # hf_hub_download returns the path it saved to; rename to the
            # local filename expected by kokoro_engine.py if different.
            if path and os.path.abspath(path) != os.path.abspath(dest):
                import shutil as _shutil
                _shutil.move(path, dest)
            logger.info("Kokoro model saved: %s", dest)
        except Exception as exc:
            logger.error("Kokoro model download failed (%s): %s", local_name, exc)
            return False

    if on_progress:
        on_progress("Kokoro models ready", 1.0)
    return is_kokoro_ready()





# ---------------------------------------------------------------------------
# Text normalization
# ---------------------------------------------------------------------------

# Inline tags some engines historically understood. Stripped for all current
# engines (Kokoro, VoxCPM, OmniVoice) since none of them parse these grammars.
_EMOTION_TAGS = re.compile(
    r'\((?:excited|happy|sad|angry|surprised|confused|nervous|confident|'
    r'satisfied|fearful|whisper|laughing|crying|shouting|serious|gentle)\)',
    re.IGNORECASE,
)
_BRACKET_TAGS = re.compile(r'\[[^\]]{1,40}\]')


def normalize_text(text: str, engine: str = "kokoro") -> str:
    """
    Normalize text before TTS to improve pronunciation and quality.

    engine: "kokoro" | "voxcpm_05b" | "voxcpm_2b" | "omnivoice"

    - Strips unknown/broken markup (all current engines read plain text)
    - Converts numbers to words (requires num2words)
    - Expands common abbreviations
    - Strips URLs
    """
    text = _BRACKET_TAGS.sub(' ', text)
    text = _EMOTION_TAGS.sub(' ', text)

    # 1. Strip URLs
    text = re.sub(r'https?://\S+|www\.\S+', '', text)

    # 2. Common abbreviations — order matters (longer matches first)
    abbrevs = [
        (r'\bDr\.(?=\s)', 'Doctor'),
        (r'\bMr\.(?=\s)', 'Mister'),
        (r'\bMrs\.(?=\s)', 'Missus'),
        (r'\bMs\.(?=\s)', 'Miss'),
        (r'\bProf\.(?=\s)', 'Professor'),
        (r'\bSgt\.(?=\s)', 'Sergeant'),
        (r'\bCpt\.(?=\s)', 'Captain'),
        (r'\bLt\.(?=\s)', 'Lieutenant'),
        (r'\bSt\.(?=\s)', 'Saint'),
        (r'\bAve\.', 'Avenue'),
        (r'\bBlvd\.', 'Boulevard'),
        (r'\bvs\.', 'versus'),
        (r'\betc\.', 'etcetera'),
        (r'\be\.g\.', 'for example'),
        (r'\bi\.e\.', 'that is'),
    ]
    for pattern, replacement in abbrevs:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # 3. Numbers to words
    try:
        from num2words import num2words

        def _replace_number(m):
            raw = m.group(0).replace(',', '')
            try:
                if '.' in raw:
                    return num2words(float(raw))
                return num2words(int(raw))
            except Exception:
                return m.group(0)

        # Match numbers with optional thousands-commas (e.g. 1,234 or 3.14)
        text = re.sub(r'\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\b|\b\d+(?:\.\d+)?\b',
                      _replace_number, text)
    except ImportError:
        pass  # num2words not installed; skip silently

    # 4. Collapse whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ---------------------------------------------------------------------------
# Reference audio pre-processing
# ---------------------------------------------------------------------------

def preprocess_reference_audio(wav_path: str, denoise: bool = True) -> str:
    """
    Clean up a reference audio clip for better voice cloning quality.

    Steps applied:
      1. Convert to mono 16-bit PCM
      2. Trim leading/trailing silence (threshold -40 dBFS, 150 ms padding)
      3. High-pass filter at 80 Hz (removes rumble, HVAC, mic handling noise)
      4. Denoise via noisereduce — stationary noise/hiss reduction (if installed)
      5. RMS normalize to -18 dBFS (consistent perceived loudness)
      6. Peak normalize to -1 dBFS (safety ceiling)

    Returns the path to a processed temp WAV (caller should delete when done).
    If pydub is unavailable the original path is returned unchanged.
    """
    import tempfile
    import os

    try:
        from pydub import AudioSegment
        from pydub.silence import detect_leading_silence
        from pydub.effects import normalize
    except ImportError:
        logger.warning("pydub not available; skipping reference audio preprocessing.")
        return wav_path

    try:
        audio = AudioSegment.from_file(wav_path)

        # Mono
        if audio.channels > 1:
            audio = audio.set_channels(1)

        # Trim silence
        start_trim = detect_leading_silence(audio, silence_threshold=-40)
        end_trim   = detect_leading_silence(audio.reverse(), silence_threshold=-40)
        duration   = len(audio)
        if start_trim + end_trim < duration:
            pad = 150  # ms of silence to keep at each end
            audio = audio[max(0, start_trim - pad): duration - max(0, end_trim - pad)]

        # High-pass filter — remove low-frequency rumble (HVAC, mic handling) below 80 Hz
        try:
            import numpy as np
            from scipy.signal import butter, sosfilt
            samples = np.array(audio.get_array_of_samples(), dtype=np.float32)
            sr = audio.frame_rate
            sos = butter(4, 80.0 / (sr / 2), btype="high", output="sos")
            filtered = sosfilt(sos, samples)
            audio = AudioSegment(
                np.clip(filtered, -32768, 32767).astype(np.int16).tobytes(),
                frame_rate=sr, sample_width=2, channels=1,
            )
        except Exception as exc:
            logger.warning("High-pass filter failed (non-fatal): %s", exc)

        # Denoise — remove stationary background noise/hiss
        if denoise:
            try:
                import numpy as np
                import noisereduce as nr

                samples = np.array(audio.get_array_of_samples(), dtype=np.float32)
                sr = audio.frame_rate
                reduced = nr.reduce_noise(y=samples, sr=sr, stationary=True, prop_decrease=0.75)
                reduced_int16 = np.clip(reduced, -32768, 32767).astype(np.int16)
                audio = AudioSegment(
                    reduced_int16.tobytes(),
                    frame_rate=sr,
                    sample_width=2,
                    channels=1,
                )
            except ImportError:
                logger.warning("noisereduce not installed — skipping denoising. "
                               "Run: pip install noisereduce")
            except Exception as exc:
                logger.warning("Denoising failed (non-fatal): %s", exc)

        # RMS normalize to -18 dBFS for consistent perceived loudness
        try:
            import numpy as np
            samples = np.array(audio.get_array_of_samples(), dtype=np.float32)
            rms = np.sqrt(np.mean(samples ** 2))
            if rms > 0:
                target_rms = 32768 * 10 ** (-18 / 20)  # -18 dBFS in int16 scale
                gain = target_rms / rms
                # Cap gain to avoid over-amplifying very quiet recordings
                gain = min(gain, 4.0)
                samples = np.clip(samples * gain, -32768, 32767).astype(np.int16)
                audio = AudioSegment(
                    samples.tobytes(),
                    frame_rate=audio.frame_rate, sample_width=2, channels=1,
                )
        except Exception as exc:
            logger.warning("RMS normalization failed (non-fatal): %s", exc)

        # Peak normalize as final safety pass
        audio = normalize(audio)

        # Write to temp file
        tmp_path = tempfile.mktemp(suffix="_ref_processed.wav")
        audio.export(tmp_path, format="wav")
        logger.info("Reference audio preprocessed: %s -> %s", wav_path, tmp_path)
        return tmp_path

    except Exception as exc:
        logger.warning("Reference audio preprocessing failed (non-fatal): %s", exc)
        return wav_path  # Fall back to original


# ---------------------------------------------------------------------------
# File readers
# ---------------------------------------------------------------------------

def read_txt(path: str) -> str:
    """Read a plain text file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def read_pdf(path: str) -> str:
    """Extract text from a PDF using pdfplumber."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


def read_docx(path: str) -> str:
    """Extract text from a DOCX file."""
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def read_epub(path: str) -> str:
    """
    Extract text from an EPUB file.

    Reads chapters in spine order, strips HTML tags, and joins paragraphs
    with double newlines so the sentence splitter sees proper boundaries.
    """
    import html
    import re
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        """Minimal HTML-to-text converter that preserves paragraph breaks."""
        BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
                      "li", "tr", "br", "blockquote", "section", "article"}

        def __init__(self):
            super().__init__()
            self._parts: list = []
            self._current = []
            self._skip = False

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style", "head"):
                self._skip = True
            if tag in self.BLOCK_TAGS and self._current:
                self._flush()

        def handle_endtag(self, tag):
            if tag in ("script", "style", "head"):
                self._skip = False
                self._current.clear()
            if tag in self.BLOCK_TAGS:
                self._flush()

        def handle_data(self, data):
            if not self._skip:
                data = data.strip()
                if data:
                    self._current.append(data)

        def handle_entityref(self, name):
            self._current.append(html.unescape(f"&{name};"))

        def handle_charref(self, name):
            self._current.append(html.unescape(f"&#{name};"))

        def _flush(self):
            text = " ".join(self._current).strip()
            if text:
                self._parts.append(text)
            self._current.clear()

        def get_text(self) -> str:
            self._flush()
            return "\n\n".join(self._parts)

    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        raise ImportError(
            "ebooklib is required for EPUB support.\n"
            "Install it with: pip install ebooklib"
        )

    book = epub.read_epub(path, options={"ignore_ncx": True})
    chapters: list = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        content = item.get_content()
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        extractor = _TextExtractor()
        extractor.feed(content)
        text = extractor.get_text().strip()
        if text:
            chapters.append(text)

    if not chapters:
        raise ValueError("No readable text found in EPUB.")

    return "\n\n".join(chapters)


def read_file(path: str) -> str:

    """
    Dispatch to the correct reader based on file extension.
    Supported: .txt, .pdf, .docx, .epub
    """
    ext = os.path.splitext(path)[1].lower()
    readers = {
        ".txt":  read_txt,
        ".pdf":  read_pdf,
        ".docx": read_docx,
        ".epub": read_epub,
    }
    reader = readers.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return reader(path)



# ---------------------------------------------------------------------------
# Audio export
# ---------------------------------------------------------------------------

def export_mp3(wav_path: str, out_path: str) -> str:
    """Convert a WAV file to MP3 using pydub + ffmpeg."""
    from pydub import AudioSegment

    if not is_ffmpeg_available():
        raise RuntimeError("ffmpeg is not available. Cannot export to MP3.")
    setup_ffmpeg()
    audio = AudioSegment.from_wav(wav_path)
    audio.export(out_path, format="mp3", bitrate="192k")
    logger.info("Exported MP3: %s", out_path)
    return out_path


def apply_speed(wav_path: str, speed: float, out_path: str = None) -> str:
    """
    Adjust playback speed of a WAV file.
    speed < 1.0 = slower, speed > 1.0 = faster.
    """
    from pydub import AudioSegment

    if abs(speed - 1.0) < 0.01:
        return wav_path

    if out_path is None:
        base, ext = os.path.splitext(wav_path)
        out_path = f"{base}_speed{ext}"

    audio = AudioSegment.from_wav(wav_path)
    # Change frame rate for speed adjustment, then restore sample rate
    adjusted = audio._spawn(
        audio.raw_data,
        overrides={"frame_rate": int(audio.frame_rate * speed)}
    ).set_frame_rate(audio.frame_rate)
    adjusted.export(out_path, format="wav")
    return out_path


def apply_volume(wav_path: str, volume_db: float, out_path: str = None) -> str:
    """Adjust volume of a WAV file by the given dB amount."""
    from pydub import AudioSegment

    if abs(volume_db) < 0.1:
        return wav_path

    if out_path is None:
        base, ext = os.path.splitext(wav_path)
        out_path = f"{base}_vol{ext}"

    audio = AudioSegment.from_wav(wav_path)
    adjusted = audio + volume_db
    adjusted.export(out_path, format="wav")
    return out_path


# ---------------------------------------------------------------------------
# Document export
# ---------------------------------------------------------------------------

def export_txt(text: str, path: str) -> str:
    """Save text to a .txt file."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    logger.info("Exported TXT: %s", path)
    return path


def export_docx(text: str, path: str) -> str:
    """Save text to a .docx file."""
    doc = DocxDocument()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    doc.save(path)
    logger.info("Exported DOCX: %s", path)
    return path


def export_pdf(text: str, path: str) -> str:
    """Save text to a .pdf file using ReportLab."""
    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    for para in text.split("\n"):
        if para.strip():
            story.append(Paragraph(para, styles["Normal"]))
            story.append(Spacer(1, 6))
    if not story:
        story.append(Paragraph("(empty)", styles["Normal"]))
    doc.build(story)
    logger.info("Exported PDF: %s", path)
    return path


def export_epub(text: str, path: str, title: str = "Document") -> str:
    """Save text to an EPUB file using ebooklib."""
    try:
        from ebooklib import epub
    except ImportError:
        raise ImportError("ebooklib is required for EPUB export.\nInstall it with: pip install ebooklib")

    book = epub.EpubBook()
    book.set_title(title)
    book.set_language("en")

    # Split on double newlines to create paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [text.strip() or "(empty)"]

    html_body = "\n".join(f"<p>{p}</p>" for p in paragraphs)
    chapter = epub.EpubHtml(title=title, file_name="content.xhtml", lang="en")
    chapter.content = (
        "<?xml version='1.0' encoding='utf-8'?>"
        "<html xmlns='http://www.w3.org/1999/xhtml'>"
        f"<head><title>{title}</title></head>"
        f"<body>{html_body}</body></html>"
    )

    book.add_item(chapter)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]

    epub.write_epub(path, book)
    logger.info("Exported EPUB: %s", path)
    return path


# ---------------------------------------------------------------------------
# System monitoring
# ---------------------------------------------------------------------------

def get_cpu_usage() -> dict:
    """
    Return CPU usage info.

    Returns dict with keys:
      - percent: float — system-wide CPU usage percentage
      - cores: int — logical CPU count
      - threads_torch: int — threads currently set for PyTorch (0 = auto)
    """
    percent = psutil.cpu_percent(interval=None)
    cores = psutil.cpu_count(logical=True) or 1
    try:
        import torch
        threads_torch = torch.get_num_threads()
    except Exception:
        threads_torch = 0
    return {
        "percent": percent,
        "cores": cores,
        "threads_torch": threads_torch,
    }


def get_ram_usage() -> dict:
    """
    Return RAM usage info for the current process and system.

    Returns dict with keys:
      - process_mb: float — RAM used by this process
      - system_percent: float — system-wide RAM usage percentage
      - system_used_gb: float — system RAM used in GB
      - system_total_gb: float — total system RAM in GB
    """
    proc = psutil.Process(os.getpid())
    mem_info = proc.memory_info()
    sys_mem = psutil.virtual_memory()
    return {
        "process_mb": mem_info.rss / (1024 * 1024),
        "system_percent": sys_mem.percent,
        "system_used_gb": sys_mem.used / (1024 ** 3),
        "system_total_gb": sys_mem.total / (1024 ** 3),
    }

def get_vram_usage() -> dict:
    """
    Return system-wide VRAM usage if a CUDA GPU is available.
    
    Returns array or None:
      - used_gb: float — system-wide GPU VRAM used in GB
      - total_gb: float — GPU VRAM total in GB
      - percent: float — VRAM usage percentage
    """
    try:
        import torch
        if torch.cuda.is_available():
            # torch.cuda.mem_get_info() returns (free, total) for the system
            free, total = torch.cuda.mem_get_info()
            used = total - free
            return {
                "used_gb": used / (1024**3),
                "total_gb": total / (1024**3),
                "percent": (used / total) * 100 if total > 0 else 0.0,
            }
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# VoxCPM / OmniVoice install-on-demand helpers
# ---------------------------------------------------------------------------

# VoxCPM variants — each maps to a HuggingFace repo and its native sample rate.
VOXCPM_VARIANTS = {
    "0.5B": {"repo": "openbmb/VoxCPM-0.5B", "sample_rate": 16000, "display": "VoxCPM 0.5B (16 kHz, light)"},
    "2B":   {"repo": "openbmb/VoxCPM2",     "sample_rate": 48000, "display": "VoxCPM 2B (48 kHz, top quality)"},
}
OMNIVOICE_REPO = "k2-fsa/OmniVoice"


def is_package_installed(pkg: str) -> bool:
    """Return True if the given pip package is importable (no side effects)."""
    import importlib.util
    try:
        return importlib.util.find_spec(pkg) is not None
    except Exception as exc:
        logger.warning("is_package_installed(%s) failed: %s", pkg, exc)
        return False


def install_package(pkg: str, on_progress=None) -> bool:
    """
    Install a pip package as a subprocess using the current Python interpreter.

    Streams stdout lines to on_progress(line, None) if provided. Progress fraction
    is None because pip output isn't quantifiable.

    Returns True on exit code 0, False otherwise.
    """
    cmd = [sys.executable, "-m", "pip", "install", pkg]
    logger.info("Installing package: %s", " ".join(cmd))
    try:
        creationflags = 0
        if sys.platform == "win32":
            creationflags = 0x08000000  # CREATE_NO_WINDOW
        proc = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            creationflags=creationflags,
        )
        for line in proc.stdout:
            line = line.rstrip()
            if on_progress:
                on_progress(line, None)
        proc.wait()
        if proc.returncode == 0:
            logger.info("Package %s installed successfully.", pkg)
            return True
        logger.error("Package %s install failed (code %d).", pkg, proc.returncode)
        return False
    except Exception as exc:
        logger.error("install_package(%s) failed: %s", pkg, exc)
        return False


def is_voxcpm_installed() -> bool:
    """Return True if the voxcpm pip package is importable."""
    return is_package_installed("voxcpm")


def is_omnivoice_installed() -> bool:
    """Return True if the omnivoice pip package is importable."""
    return is_package_installed("omnivoice")


def is_triton_installed() -> bool:
    """Return True if triton (needed for torch.compile) is importable."""
    return is_package_installed("triton")


def ensure_triton(on_progress=None) -> bool:
    """
    Install triton so torch.compile can fuse kernels (3-10x faster inference
    on VoxCPM/OmniVoice). Non-fatal: returns False if install fails, but the
    caller should continue — the models still run, just slower.

    On Windows the wheel is `triton-windows`; elsewhere it's plain `triton`.
    """
    if is_triton_installed():
        return True
    pkg = "triton-windows" if sys.platform == "win32" else "triton"
    if on_progress:
        on_progress(f"Installing {pkg} for faster inference...", None)
    ok = install_package(pkg, on_progress=on_progress)
    if not ok:
        logger.warning(
            "triton install failed — models will still work but run much "
            "slower. You can install manually later: pip install %s", pkg,
        )
    return ok


def hf_cache_has_model(repo_id: str) -> bool:
    """
    Return True if HuggingFace hub has the given repo cached locally.

    Tries huggingface_hub.try_to_load_from_cache first; falls back to checking
    for the ~/.cache/huggingface/hub/models--<org>--<name> directory so the
    check still works before huggingface_hub is installed.
    """
    try:
        from huggingface_hub import try_to_load_from_cache
        from huggingface_hub.constants import HF_HUB_CACHE
        # try_to_load_from_cache returns a path, None (not cached), or _CACHED_NO_EXIST
        for candidate in ("config.json", "model.safetensors", "pytorch_model.bin"):
            result = try_to_load_from_cache(repo_id=repo_id, filename=candidate)
            if isinstance(result, str) and os.path.isfile(result):
                return True
        # Fall through to filesystem check using HF cache root
        cache_root = HF_HUB_CACHE
    except Exception:
        cache_root = os.path.join(os.path.expanduser("~"), ".cache", "huggingface", "hub")

    try:
        dir_name = "models--" + repo_id.replace("/", "--")
        target = os.path.join(cache_root, dir_name)
        return os.path.isdir(target)
    except Exception as exc:
        logger.warning("hf_cache_has_model(%s) fallback failed: %s", repo_id, exc)
        return False


def is_voxcpm_ready(variant: str) -> bool:
    """Return True if voxcpm is installed AND the variant's weights are cached."""
    cfg = VOXCPM_VARIANTS.get(variant)
    if not cfg:
        logger.warning("is_voxcpm_ready: unknown variant %r", variant)
        return False
    return is_voxcpm_installed() and hf_cache_has_model(cfg["repo"])


def is_omnivoice_ready() -> bool:
    """Return True if omnivoice is installed AND its weights are cached."""
    return is_omnivoice_installed() and hf_cache_has_model(OMNIVOICE_REPO)


def setup_voxcpm(variant: str, on_progress=None) -> bool:
    """
    Ensure the voxcpm package is installed and the variant weights are downloaded.

    variant: "0.5B" or "2B"

    Steps:
      1. pip install voxcpm (if not installed)
      2. Instantiate VoxCPM.from_pretrained(repo, load_denoiser=False) to trigger
         HuggingFace cache download, then drop the instance.

    on_progress(message, fraction) is called at key stages.
    Returns True on full success.
    """
    cfg = VOXCPM_VARIANTS.get(variant)
    if not cfg:
        logger.error("setup_voxcpm: unknown variant %r", variant)
        return False
    repo = cfg["repo"]

    if not is_voxcpm_installed():
        if on_progress:
            on_progress("Installing voxcpm package...", None)
        if not install_package("voxcpm", on_progress=on_progress):
            logger.error("setup_voxcpm: voxcpm package install failed.")
            return False

    # Best-effort triton install for torch.compile (big speedup, non-fatal).
    ensure_triton(on_progress=on_progress)

    if on_progress:
        on_progress("Downloading model weights...", None)
    try:
        from huggingface_hub import snapshot_download
        logger.info("Downloading VoxCPM weights: %s", repo)
        snapshot_download(
            repo_id=repo,
            allow_patterns=["*.json", "*.safetensors", "*.pth", "*.bin", "*.model", "*.txt", "tokenizer*"],
        )
    except Exception as exc:
        logger.error("setup_voxcpm(%s) weight download failed: %s", variant, exc)
        return False

    if on_progress:
        on_progress("Ready", 1.0)
    logger.info("VoxCPM %s ready.", variant)
    return True


def setup_omnivoice(on_progress=None) -> bool:
    """
    Ensure the omnivoice package is installed and its weights are downloaded.

    Steps:
      1. pip install omnivoice (if not installed)
      2. Instantiate OmniVoice.from_pretrained(OMNIVOICE_REPO, device_map="cpu")
         to trigger HuggingFace cache download (CPU avoids CUDA alloc), then
         drop the instance.

    on_progress(message, fraction) is called at key stages.
    Returns True on full success.
    """
    if not is_omnivoice_installed():
        if on_progress:
            on_progress("Installing omnivoice package...", None)
        if not install_package("omnivoice", on_progress=on_progress):
            logger.error("setup_omnivoice: omnivoice package install failed.")
            return False

    # Best-effort triton install for torch.compile (big speedup, non-fatal).
    ensure_triton(on_progress=on_progress)

    if on_progress:
        on_progress("Downloading model weights...", None)
    try:
        from huggingface_hub import snapshot_download
        logger.info("Downloading OmniVoice weights: %s", OMNIVOICE_REPO)
        snapshot_download(
            repo_id=OMNIVOICE_REPO,
            allow_patterns=["*.json", "*.safetensors", "*.pth", "*.bin", "*.model", "*.txt", "tokenizer*"],
        )
    except Exception as exc:
        logger.error("setup_omnivoice weight download failed: %s", exc)
        return False

    if on_progress:
        on_progress("Ready", 1.0)
    logger.info("OmniVoice ready.")
    return True
